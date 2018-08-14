from docx import Document
docStr=Document('./doc/1.docx')
for paragraph in docStr.paragraphs:
    parStr=paragraph.text
    print(parStr)

for table in docStr.tables:
    for row in table.rows:
            print(row)

